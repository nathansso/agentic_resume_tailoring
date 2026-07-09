"""
Resume Formatter Agent — Converts tailored JSON output into a formatted resume.

Primary output path: LaTeX → PDF via pdflatex (Jake's Resume layout).
Also supports .tex source export and .docx export.
Markdown output is retained for debugging but is not the intended pipeline.
"""
import logging
import re
import shutil
import warnings
from typing import Dict, List, Optional
from uuid import UUID

from sqlmodel import Session, select

from database.db import engine
from database.models import Education, Skill, User, UserSkill
from agents.skill_postprocessor import normalize_skill_name, should_reject_skill

logger = logging.getLogger(__name__)

_DEFAULT_SECTION_ORDER = ["education", "experience", "projects", "skills"]

# ── Jake's Resume preamble (exact match to https://github.com/jakegut/resume) ──

_JAKE_PREAMBLE = r"""%-------------------------
% Resume in Latex
% Based off of: https://github.com/jakegut/resume
% License : MIT
%------------------------

\documentclass[letterpaper,11pt]{article}

\usepackage{latexsym}
\usepackage[empty]{fullpage}
\usepackage{titlesec}
\usepackage{marvosym}
\usepackage[usenames,dvipsnames]{color}
\usepackage{verbatim}
\usepackage{enumitem}
\usepackage[hidelinks]{hyperref}
\usepackage{fancyhdr}
\usepackage[english]{babel}
\usepackage{tabularx}
\ifdefined\pdftexversion
  \usepackage[utf8]{inputenc}
  \input{glyphtounicode}
\fi

\pagestyle{fancy}
\fancyhf{}
\fancyfoot{}
\renewcommand{\headrulewidth}{0pt}
\renewcommand{\footrulewidth}{0pt}

\addtolength{\oddsidemargin}{-0.5in}
\addtolength{\evensidemargin}{-0.5in}
\addtolength{\textwidth}{1in}
\addtolength{\topmargin}{-.5in}
\addtolength{\textheight}{1.0in}

\urlstyle{same}

\raggedbottom
\raggedright
\setlength{\tabcolsep}{0in}

\titleformat{\section}{
  \vspace{-4pt}\scshape\raggedright\large
}{}{0em}{}[\color{black}\titlerule \vspace{-5pt}]

\ifdefined\pdftexversion\pdfgentounicode=1\fi

%-------------------------
% Custom commands
\newcommand{\resumeItem}[1]{
  \item\small{
    {#1 \vspace{-2pt}}
  }
}

\newcommand{\resumeSubheading}[4]{
  \vspace{-2pt}\item
    \begin{tabular*}{0.97\textwidth}[t]{l@{\extracolsep{\fill}}r}
      \textbf{#1} & #2 \\
      \textit{\small#3} & \textit{\small #4} \\
    \end{tabular*}\vspace{-7pt}
}

\newcommand{\resumeSubSubheading}[2]{
    \item
    \begin{tabular*}{0.97\textwidth}{l@{\extracolsep{\fill}}r}
      \textit{\small#1} & \textit{\small #2} \\
    \end{tabular*}\vspace{-7pt}
}

\newcommand{\resumeProjectHeading}[2]{
    \item
    \begin{tabular*}{0.97\textwidth}{l@{\extracolsep{\fill}}r}
      \small#1 & #2 \\
    \end{tabular*}\vspace{-7pt}
}

\newcommand{\resumeSubItem}[1]{\resumeItem{#1}\vspace{-4pt}}

\renewcommand\labelitemii{$\vcenter{\hbox{\tiny$\bullet$}}$}

\newcommand{\resumeSubHeadingListStart}{\begin{itemize}[leftmargin=0.15in, label={}]}
\newcommand{\resumeSubHeadingListEnd}{\end{itemize}}
\newcommand{\resumeItemListStart}{\begin{itemize}}
\newcommand{\resumeItemListEnd}{\end{itemize}\vspace{-5pt}}
"""

# ── TeX escaping ──────────────────────────────────────────────────────────────

_TEX_SPECIAL = {
    "\\": r"\textbackslash{}",
    "&":  r"\&",
    "%":  r"\%",
    "$":  r"\$",
    "#":  r"\#",
    "_":  r"\_",
    "{":  r"\{",
    "}":  r"\}",
    "~":  r"\textasciitilde{}",
    "^":  r"\textasciicircum{}",
}
_TEX_ESCAPE_RE = re.compile(r"[\\&%$#_{}~^]")


def _escape_tex(text: str) -> str:
    """Escape special LaTeX characters in user-supplied text (single-pass)."""
    return _TEX_ESCAPE_RE.sub(lambda m: _TEX_SPECIAL[m.group()], text)


_MD_LINK_RE = re.compile(r"\[[^\]]+\]\([^)]+\)")


def _convert_inline(text: str) -> str:
    """Convert **bold** / *italic* / [text](url) markdown to LaTeX, escaping surrounding text."""
    parts = re.split(r"(\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*|\*[^*]+\*)", text)
    result = []
    for part in parts:
        link = _MD_LINK_RE.match(part)
        if link:
            label, url = part[1:].split("](", 1)
            url = url[:-1]
            result.append(
                rf"\href{{{_escape_tex(url)}}}{{\underline{{{_escape_tex(label)}}}}}"
            )
        elif part.startswith("**") and part.endswith("**"):
            result.append(r"\textbf{" + _escape_tex(part[2:-2]) + "}")
        elif part.startswith("*") and part.endswith("*"):
            result.append(r"\textit{" + _escape_tex(part[1:-1]) + "}")
        else:
            result.append(_escape_tex(part))
    return "".join(result)


# ── PDF compilation ───────────────────────────────────────────────────────────

def _compile_tex_to_pdf(tex_str: str) -> bytes:
    """Compile tex_str to PDF bytes. Prefers tectonic, falls back to pdflatex."""
    import os
    import subprocess
    import tempfile

    # Resolve engine: check PATH first, then check alongside this Python executable
    import sys as _sys
    _py_bin = os.path.dirname(_sys.executable)
    tectonic = (
        shutil.which("tectonic")
        or (os.path.join(_py_bin, "tectonic.exe") if os.path.exists(os.path.join(_py_bin, "tectonic.exe")) else None)
        or (os.path.join(_py_bin, "tectonic") if os.path.exists(os.path.join(_py_bin, "tectonic")) else None)
    )
    pdflatex = shutil.which("pdflatex")
    if not tectonic and not pdflatex:
        raise RuntimeError("No LaTeX engine found. Install tectonic or pdflatex.")

    with tempfile.TemporaryDirectory() as tmpdir:
        tex_path = os.path.join(tmpdir, "resume.tex")
        pdf_path = os.path.join(tmpdir, "resume.pdf")
        with open(tex_path, "w", encoding="utf-8") as f:
            f.write(tex_str)

        if tectonic:
            result = subprocess.run(
                [tectonic, "--outdir", tmpdir, tex_path],
                capture_output=True,
                timeout=120,
            )
        else:
            result = subprocess.run(
                [pdflatex, "-interaction=nonstopmode", "-output-directory", tmpdir, tex_path],
                capture_output=True,
                timeout=60,
            )

        if not os.path.exists(pdf_path):
            out = (result.stdout + result.stderr).decode("utf-8", errors="replace")
            engine = "tectonic" if tectonic else "pdflatex"
            raise RuntimeError(f"{engine} failed (exit {result.returncode}):\n{out[-3000:]}")

        with open(pdf_path, "rb") as f:
            return f.read()


# ── One-page fit enforcement (issue #34) ──────────────────────────────────────


def _pdf_page_count(pdf_bytes: bytes) -> int:
    """Number of pages in a PDF byte string, via pypdf."""
    import io

    from pypdf import PdfReader

    return len(PdfReader(io.BytesIO(pdf_bytes)).pages)


# One-page trim floors (issue #72). We shave a project/experience down to this
# many bullets before touching anything else, so descriptions stay substantive
# instead of every entry collapsing to a single line.
_PROJECT_BULLET_FLOOR = 2
_EXP_BULLET_FLOOR = 2


def _trim_one_bullet(content: Dict) -> Optional[Dict]:
    """
    Return a copy of `content` with exactly one reduction applied, or None when
    nothing more can be trimmed (issue #34). Text-only reduction — never touches
    the preamble, geometry, font, section list, or skills.

    Ladder (issue #72 — favor fewer, well-described items over many starved
    ones, and protect experiences more than projects):
      1. Trim the lowest-ranked project's last bullet, but only down to
         _PROJECT_BULLET_FLOOR. Projects are pre-ordered descending by relevance
         (issue #47), so the last project is the least relevant.
      2. Trim the lowest-ranked experience's last bullet, down to _EXP_BULLET_FLOOR.
      3. Everything is at the floor and it still overflows: drop the entire
         lowest-ranked project (keeping at least one) rather than starving the
         survivors below the floor.
      4. Last resorts for project-light / single-project resumes: shave below the
         floor — experiences first (projects are the showcase), then the last
         remaining project — never below one bullet.
    """
    import copy

    trimmed = copy.deepcopy(content)
    projects = trimmed.get("projects") if isinstance(trimmed.get("projects"), list) else []
    experiences = trimmed.get("experiences") if isinstance(trimmed.get("experiences"), list) else []

    def _shave(items, floor):
        for item in reversed(items):
            bullets = item.get("bullets")
            if isinstance(bullets, list) and len(bullets) > floor:
                bullets.pop()
                return True
        return False

    # 1–2. Trim fat down to the floor.
    if _shave(projects, _PROJECT_BULLET_FLOOR):
        return trimmed
    if _shave(experiences, _EXP_BULLET_FLOOR):
        return trimmed

    # 3. Drop the weakest whole project before starving the rest.
    if len(projects) > 1:
        projects.pop()
        return trimmed

    # 4. Below-floor last resorts, protecting projects over experiences.
    if _shave(experiences, 1):
        return trimmed
    if _shave(projects, 1):
        return trimmed

    return None


def _fit_to_one_page(content: Dict, render_fn, page_count_fn, max_iters: int = 30) -> Dict:
    """
    Deterministically shrink `content` until it renders to one page (issue #34).

    `render_fn(content) -> pdf_bytes` and `page_count_fn(pdf_bytes) -> int` are
    injected so the loop is testable without a real LaTeX engine. Returns the final
    content dict (the caller renders it once more, or reuses the last render). Stops
    when the content fits, the reducer is exhausted, or `max_iters` is reached.
    """
    current = content
    for _ in range(max_iters):
        if page_count_fn(render_fn(current)) <= 1:
            return current
        reduced = _trim_one_bullet(current)
        if reduced is None:
            return current
        current = reduced
    return current


# Date strings that mean "no real date" — rendered as an empty date rather than
# printing the literal placeholder (issue #72). Mirrors tailor._PLACEHOLDER_DATE_TOKENS
# so exports render cleanly even for content built directly from the DB.
_PLACEHOLDER_DATE_TOKENS = {
    "", "not specified", "unknown", "unspecified", "n/a", "na", "none", "tbd", "-",
}


def _clean_date(value) -> str:
    """Blank out placeholder date strings; pass real dates through unchanged."""
    v = str(value or "").strip()
    return "" if v.lower() in _PLACEHOLDER_DATE_TOKENS else v


def sanitize_text(text: str) -> str:
    """Replace common Unicode characters with ASCII-safe equivalents."""
    for char, rep in {
        "–": "-", "—": "-", "'": "'", "'": "'",
        "“": '"', "”": '"', "…": "...",
        " ": " ", "​": "",
    }.items():
        text = text.replace(char, rep)
    return text


# ── Formatter ─────────────────────────────────────────────────────────────────

class ResumeFormatterAgent:
    """Formats tailored content into PDF (LaTeX), .tex source, .docx, or Markdown."""

    def __init__(self, user_id: UUID):
        self.user_id = user_id
        self._style: Optional[dict] = self._load_style()

    def _load_style(self) -> Optional[dict]:
        try:
            from services import get_resume_style
            return get_resume_style(self.user_id)
        except Exception:
            return None

    def _resolve_label(self, key: str) -> str:
        return (self._style or {}).get("section_labels", {}).get(key, key.capitalize())

    def _get_education(self) -> List[Education]:
        """This user's education rows in insertion (resume document) order."""
        with Session(engine) as session:
            return list(session.exec(
                select(Education)
                .where(Education.user_id == self.user_id)
                .order_by(Education.created_at)
            ).all())

    @staticmethod
    def _education_degree_line(entry: Education) -> str:
        return f"{entry.degree}, GPA: {entry.gpa}" if entry.gpa else entry.degree

    @staticmethod
    def _education_date_line(entry: Education) -> str:
        if entry.start_date and entry.end_date:
            return f"{entry.start_date} -- {entry.end_date}"
        return entry.end_date or entry.start_date or ""

    def _build_known_labels(self) -> frozenset:
        labels: set = set()
        for v in (self._style or {}).get("section_labels", {}).values():
            labels.add(v)
        return frozenset(labels)

    # ── .tex source ───────────────────────────────────────────────────────────

    def format_tex(
        self,
        tailored_content: Dict,
        job_title: str = "",
        section_order: Optional[List[str]] = None,
    ) -> str:
        """Return the raw LaTeX source string (Jake's Resume layout)."""
        return self._build_tex(tailored_content, section_order=section_order)

    def fit_content_to_one_page(
        self,
        tailored_content: Dict,
        section_order: Optional[List[str]] = None,
    ) -> Dict:
        """
        Trim `tailored_content` until it renders to a single page, so the stored
        content (and every output derived from it — .tex source, live preview,
        exports) agrees on a one-page layout. Returns the content unchanged when
        it already fits or when no LaTeX engine is available (offline/test
        environments) — enforcement then falls back to format_pdf's own guard.
        """
        def render(content: Dict) -> bytes:
            return _compile_tex_to_pdf(
                self._build_tex(content, section_order=section_order)
            )

        try:
            if _pdf_page_count(render(tailored_content)) <= 1:
                return tailored_content
            return _fit_to_one_page(tailored_content, render, _pdf_page_count)
        except Exception:
            return tailored_content

    # ── PDF ───────────────────────────────────────────────────────────────────

    def format_pdf(
        self,
        tailored_content: Dict,
        job_title: str = "",
        section_order: Optional[List[str]] = None,
    ) -> bytes:
        """
        Compile to PDF and return raw bytes, enforcing a strict one-page layout
        (issue #34). A resume that already fits compiles exactly once (fast path);
        an overflowing one has bullet text trimmed deterministically and is
        recompiled until it fits — font, margins, and spacing are never changed.
        """
        def render(content: Dict) -> bytes:
            return _compile_tex_to_pdf(
                self._build_tex(content, section_order=section_order)
            )

        pdf_bytes = render(tailored_content)
        if _pdf_page_count(pdf_bytes) <= 1:
            return pdf_bytes

        fitted = _fit_to_one_page(tailored_content, render, _pdf_page_count)
        return render(fitted)

    # ── DOCX ──────────────────────────────────────────────────────────────────

    def format_docx(
        self,
        tailored_content: Dict,
        job_title: str = "",
        section_order: Optional[List[str]] = None,
    ) -> bytes:
        """Generate a Word document (.docx) mirroring the Jake's Resume layout."""
        import io
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor

        doc = Document()

        # Tight margins matching Jake's 0.5in top/bottom, 1in sides
        for sec in doc.sections:
            sec.top_margin    = Inches(0.5)
            sec.bottom_margin = Inches(0.5)
            sec.left_margin   = Inches(0.5)
            sec.right_margin  = Inches(0.5)

        # Compact default paragraph spacing
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(11)
        doc.styles["Normal"].paragraph_format.space_before = Pt(0)
        doc.styles["Normal"].paragraph_format.space_after  = Pt(2)

        def _add_bottom_border(paragraph):
            pPr  = paragraph._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "4")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "auto")
            pBdr.append(bot)
            pPr.append(pBdr)

        def _section_heading(label: str):
            p   = doc.add_paragraph()
            run = p.add_run(label.upper())
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            _add_bottom_border(p)

        def _subheading(left_bold: str, right: str, left_italic: str = "", right_italic: str = ""):
            """Two-line entry: bold-name | right, then italic-subtitle | right-italic."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(0)
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.5), 2)   # right-align at text width
            r = p.add_run(left_bold)
            r.bold = True
            r.font.size = Pt(11)
            if right:
                p.add_run("\t" + right).font.size = Pt(10)
            if left_italic:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after  = Pt(0)
                tab_stops2 = p2.paragraph_format.tab_stops
                tab_stops2.add_tab_stop(Inches(6.5), 2)
                r2 = p2.add_run(left_italic)
                r2.italic = True
                r2.font.size = Pt(10)
                if right_italic:
                    p2.add_run("\t" + right_italic).font.size = Pt(10)

        def _bullet(text: str):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent   = Inches(0.25)
            p.paragraph_format.space_before  = Pt(0)
            p.paragraph_format.space_after   = Pt(1)
            # Strip markdown bold/italic and flatten [text](url) links to "text (url)" — add as plain text
            clean = re.sub(r"\*+([^*]+)\*+", r"\1", text.strip())
            clean = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", clean)
            p.add_run(clean).font.size = Pt(10)

        # ── Header ────────────────────────────────────────────────────────────
        with Session(engine) as session:
            user = session.exec(select(User).where(User.user_id == self.user_id)).first()

        if user:
            p_name = doc.add_paragraph()
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p_name.add_run(user.name)
            r.bold = True
            r.font.size = Pt(18)

            parts = []
            if user.phone:
                parts.append(user.phone)
            if user.email and user.email != "user@example.com":
                parts.append(user.email)
            if user.linkedin_url:
                parts.append(user.linkedin_url)
            if user.github_username:
                parts.append(f"github.com/{user.github_username}")
            if user.portfolio_url:
                parts.append(user.portfolio_url)
            if user.location:
                parts.append(user.location)

            p_contact = doc.add_paragraph(" | ".join(parts))
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p_contact.runs:
                p_contact.runs[0].font.size = Pt(9)

        style = self._style or {}
        order = (
            (section_order or style.get("section_order") or _DEFAULT_SECTION_ORDER)
        )

        seen: set = set()
        sections_to_render = []
        for key in order:
            if key not in seen:
                seen.add(key)
                sections_to_render.append(key)
        for key in _DEFAULT_SECTION_ORDER:
            if key not in seen:
                seen.add(key)
                sections_to_render.append(key)

        for key in sections_to_render:
            if key == "education":
                edu_entries = self._get_education()
                if not edu_entries:
                    continue
                _section_heading(self._resolve_label("education"))
                for entry in edu_entries:
                    _subheading(
                        entry.institution, entry.location or "",
                        self._education_degree_line(entry),
                        self._education_date_line(entry),
                    )

            elif key == "experience":
                exps = tailored_content.get("experiences", [])
                if not exps:
                    continue
                _section_heading(self._resolve_label("experience"))
                for exp in exps:
                    _subheading(
                        exp.get("title", ""),
                        f"{exp.get('start_date','')} -- {exp.get('end_date','')}".strip(" --"),
                        exp.get("company", ""),
                        exp.get("location", ""),
                    )
                    for b in exp.get("bullets", []):
                        _bullet(b)

            elif key == "projects":
                projs = tailored_content.get("projects", [])
                if not projs:
                    continue
                _section_heading(self._resolve_label("projects"))
                for proj in projs:
                    name  = proj.get("name", "")
                    techs = proj.get("tech_stack", proj.get("technologies", ""))
                    dates = proj.get("date_range", proj.get("dates", ""))
                    heading = f"{name} | {techs}" if techs else name
                    links = [u for u in (proj.get("repo_url"), proj.get("demo_url")) if u]
                    if links:
                        heading += " | " + " | ".join(links)
                    _subheading(heading, dates)
                    for b in proj.get("bullets", []):
                        _bullet(b)

            elif key == "skills":
                ranked = tailored_content.get("skills_ranked")
                cats = self._get_skill_categories(ranked)
                if not cats:
                    continue
                _section_heading("Technical Skills")
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                for cat, skills in self._ordered_skill_cats(cats, ranked):
                    r_bold = p.add_run(cat + ": ")
                    r_bold.bold = True
                    r_bold.font.size = Pt(10)
                    r_text = p.add_run(", ".join(skills) + "    ")
                    r_text.font.size = Pt(10)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── Internal: build .tex document ─────────────────────────────────────────

    def _build_tex(
        self,
        tailored_content: Dict,
        section_order: Optional[List[str]] = None,
    ) -> str:
        style = self._style or {}
        order = section_order or style.get("section_order") or _DEFAULT_SECTION_ORDER

        builders = {
            "education": self._build_tex_education,
            "experience": lambda: self._build_tex_experiences(
                tailored_content.get("experiences", [])
            ),
            "projects": lambda: self._build_tex_projects(
                tailored_content.get("projects", [])
            ),
            "skills": lambda: self._build_tex_skills(
                tailored_content.get("skills_ranked")
            ),
        }

        # Each block is prefixed with an "%% ART-SECTION: <key>" comment marker
        # (issue #71) so the web editor can reorder sections/bullets as pure
        # text-block moves — even after the user hand-edits the source.
        def _marked(key: str, block: str) -> str:
            return f"%% ART-SECTION: {key}\n{block}"

        body_parts = [_marked("header", self._build_tex_header())]
        seen: set = set()
        for key in order:
            if key in seen or key not in builders:
                continue
            seen.add(key)
            block = builders[key]()
            if block:
                body_parts.append(_marked(key, block))
        for key, fn in builders.items():
            if key not in seen:
                block = fn()
                if block:
                    body_parts.append(_marked(key, block))

        body = "\n\n".join(body_parts)
        return _JAKE_PREAMBLE + "\n\\begin{document}\n\n" + body + "\n\n\\end{document}\n"

    def _build_tex_header(self) -> str:
        """Jake's centered header: large scshape name + pipe-separated contact line."""
        with Session(engine) as session:
            user = session.exec(select(User).where(User.user_id == self.user_id)).first()
            if not user:
                return ""

        name = _escape_tex(user.name)

        parts = []
        if user.phone:
            parts.append(_escape_tex(user.phone))
        if user.email and user.email != "user@example.com":
            e = _escape_tex(user.email)
            parts.append(rf"\href{{mailto:{e}}}{{\underline{{{e}}}}}")
        if user.linkedin_url:
            u = _escape_tex(user.linkedin_url)
            display = u.replace("https://", "").replace("http://", "")
            parts.append(rf"\href{{{u}}}{{\underline{{{display}}}}}")
        if user.github_username:
            u   = _escape_tex(f"https://github.com/{user.github_username}")
            dis = f"github.com/{_escape_tex(user.github_username)}"
            parts.append(rf"\href{{{u}}}{{\underline{{{dis}}}}}")
        if user.portfolio_url:
            u = _escape_tex(user.portfolio_url)
            display = u.replace("https://", "").replace("http://", "")
            parts.append(rf"\href{{{u}}}{{\underline{{{display}}}}}")
        if user.location:
            parts.append(_escape_tex(user.location))

        # Jake splits contact onto multiple source lines with $|$ separators
        contact = " $|$ \n    ".join(parts)

        return (
            r"\begin{center}" + "\n"
            rf"    \textbf{{\Huge \scshape {name}}} \\ \vspace{{1pt}}" + "\n"
            rf"    \small {contact}" + "\n"
            r"\end{center}"
        )

    def _build_tex_education(self) -> str:
        """Jake's education section: one \\resumeSubheading per Education row.

        Omitted entirely when the user has no education rows — never fabricated.
        """
        entries = self._get_education()
        if not entries:
            return ""
        label = _escape_tex(self._resolve_label("education"))
        lines = [
            rf"\section{{{label}}}",
            r"  \resumeSubHeadingListStart",
        ]
        for entry in entries:
            lines += [
                r"    \resumeSubheading",
                f"      {{{_escape_tex(entry.institution)}}}{{{_escape_tex(entry.location or '')}}}",
                f"      {{{_escape_tex(self._education_degree_line(entry))}}}"
                f"{{{_escape_tex(self._education_date_line(entry))}}}",
            ]
        lines.append(r"  \resumeSubHeadingListEnd")
        return "\n".join(lines)

    def _build_tex_experiences(self, experiences: List[Dict]) -> str:
        """Jake's experience section: \\resumeSubheading{Title}{Dates}{Company}{Location}."""
        if not experiences:
            return ""

        label = _escape_tex(self._resolve_label("experience"))
        lines = [
            rf"\section{{{label}}}",
            r"  \resumeSubHeadingListStart",
            "",
        ]

        for exp in experiences:
            title    = _escape_tex(exp.get("title", ""))
            company  = _escape_tex(exp.get("company", ""))
            start    = _escape_tex(_clean_date(exp.get("start_date")))
            end      = _escape_tex(_clean_date(exp.get("end_date")))
            location = _escape_tex(exp.get("location", ""))
            if start and end:
                dates = f"{start} -- {end}"
            else:
                dates = end or start

            lines += [
                r"    \resumeSubheading",
                f"      {{{title}}}{{{dates}}}",
                f"      {{{company}}}{{{location}}}",
            ]
            # Skip the itemize entirely when there are no bullets — an empty
            # \resumeItemListStart/End renders as blank space / a LaTeX error (issue #72).
            bullets = [b for b in (exp.get("bullets") or []) if b and b.strip()]
            if bullets:
                lines.append(r"      \resumeItemListStart")
                for bullet in bullets:
                    lines.append(
                        r"        \resumeItem{" + _convert_inline(bullet.strip()) + "}"
                    )
                lines.append(r"      \resumeItemListEnd")
            lines.append("")

        lines.append(r"  \resumeSubHeadingListEnd")
        return "\n".join(lines)

    def _build_tex_projects(self, projects: List[Dict]) -> str:
        """Jake's projects section: \\resumeProjectHeading{\\textbf{Name} $|$ \\emph{Stack}}{Date}."""
        if not projects:
            return ""

        label = _escape_tex(self._resolve_label("projects"))
        lines = [
            rf"\section{{{label}}}",
            r"    \resumeSubHeadingListStart",
        ]

        for proj in projects:
            name  = _escape_tex(proj.get("name", "Untitled"))
            techs = _escape_tex(proj.get("tech_stack", proj.get("technologies", "")))
            dates = _escape_tex(_clean_date(proj.get("date_range", proj.get("dates", ""))))

            heading = (
                rf"\textbf{{{name}}} $|$ \emph{{{techs}}}" if techs
                else rf"\textbf{{{name}}}"
            )

            links = []
            if proj.get("repo_url"):
                links.append(rf"\href{{{_escape_tex(proj['repo_url'])}}}{{\underline{{Repo}}}}")
            if proj.get("demo_url"):
                links.append(rf"\href{{{_escape_tex(proj['demo_url'])}}}{{\underline{{Demo}}}}")
            if links:
                heading += " $|$ " + " $|$ ".join(links)

            lines += [
                r"      \resumeProjectHeading",
                f"          {{{heading}}}{{{dates}}}",
            ]
            # Skip the itemize entirely when there are no bullets (issue #72).
            bullets = [b for b in (proj.get("bullets") or []) if b and b.strip()]
            if bullets:
                lines.append(r"          \resumeItemListStart")
                for bullet in bullets:
                    lines.append(
                        r"            \resumeItem{" + _convert_inline(bullet.strip()) + "}"
                    )
                lines.append(r"          \resumeItemListEnd")

        lines.append(r"    \resumeSubHeadingListEnd")
        return "\n".join(lines)

    def _build_tex_skills(self, ranked: Optional[List[Dict]] = None) -> str:
        """Jake's Technical Skills section: \\textbf{Category}{: skill, skill, ...} \\\\."""
        cats = self._get_skill_categories(ranked)
        if not cats:
            return ""

        skill_lines = []
        for cat, skills in self._ordered_skill_cats(cats, ranked):
            skills_str = _escape_tex(", ".join(skills))
            skill_lines.append(
                rf"     \textbf{{{_escape_tex(cat)}}}{{: {skills_str}}} \\"
            )

        return "\n".join([
            r"\section{Technical Skills}",
            r" \begin{itemize}[leftmargin=0.15in, label={}]",
            r"    \small{\item{",
            *skill_lines,
            r"    }}",
            r" \end{itemize}",
        ])

    # Static fallback category order, used only when no JD-ranked list is given.
    _STATIC_CAT_ORDER = [
        "Languages & Libraries", "AI & Machine Learning",
        "Data Engineering", "Tools", "Cloud", "Other",
    ]

    def _get_skill_categories(
        self, ranked: Optional[List[Dict]] = None
    ) -> Dict[str, List[str]]:
        """
        Return {category: [skill_name, ...]}.

        When `ranked` (a JD-scored list of {name, category, score} from the
        tailoring pipeline, issue #54) is given, categories and within-category
        skills preserve that relevance order. Otherwise fall back to the user's
        full skill table (rendered later in the static order, sorted A→Z).
        """
        if ranked:
            cats: Dict[str, List[str]] = {}
            for item in ranked:
                name = item.get("name")
                if not name or should_reject_skill(name):
                    continue
                canonical = normalize_skill_name(name)
                cat = self._normalize_category(item.get("category") or "Other")
                cats.setdefault(cat, [])
                if canonical not in cats[cat]:
                    cats[cat].append(canonical)
            return cats

        with Session(engine) as session:
            user_skills = session.exec(
                select(UserSkill).where(UserSkill.user_id == self.user_id)
            ).all()
            cats = {}
            for us in user_skills:
                skill = session.exec(
                    select(Skill).where(Skill.skill_id == us.skill_id)
                ).first()
                if not skill or should_reject_skill(skill.name):
                    continue
                canonical = normalize_skill_name(skill.name)
                cat = self._normalize_category(skill.category or "Other")
                cats.setdefault(cat, [])
                if canonical not in cats[cat]:
                    cats[cat].append(canonical)
        return cats

    def _ordered_skill_cats(
        self, cats: Dict[str, List[str]], ranked: Optional[List[Dict]] = None
    ) -> List[tuple]:
        """
        Resolve final (category, [skills]) render order.

        Ranked input is already relevance-ordered (categories by first
        appearance, skills by score), so it is rendered as-is. The untailored
        fallback uses the static category order with skills sorted A→Z.
        """
        if ranked:
            return list(cats.items())
        cats = dict(cats)
        pairs: List[tuple] = []
        for cat in self._STATIC_CAT_ORDER:
            if cat in cats:
                pairs.append((cat, sorted(cats.pop(cat))))
        for cat, skills in cats.items():
            pairs.append((cat, sorted(skills)))
        return pairs

    # ── Deprecated: Markdown ──────────────────────────────────────────────────

    def format_markdown(
        self,
        tailored_content: Dict,
        job_title: str = "",
        section_order: Optional[List[str]] = None,
    ) -> str:
        """[DEPRECATED] Plain Markdown resume string.

        Use format_pdf(), format_tex(), or format_docx() instead.
        Retained for debugging and plain-text inspection only.
        """
        warnings.warn(
            "format_markdown() is deprecated — use format_pdf(), format_tex(), or format_docx().",
            DeprecationWarning,
            stacklevel=2,
        )
        style = self._style or {}
        order = section_order or style.get("section_order") or _DEFAULT_SECTION_ORDER
        bullet = style.get("bullet_prefix", "- ")

        builders = {
            "education": self._build_education,
            "experience": lambda: self._build_experiences(
                tailored_content.get("experiences", []),
                label=self._resolve_label("experience"),
                bullet_prefix=bullet,
            ),
            "projects": lambda: self._build_projects(
                tailored_content.get("projects", []),
                label=self._resolve_label("projects"),
                bullet_prefix=bullet,
            ),
            "skills": lambda: self._build_skills(
                label=self._resolve_label("skills"),
                ranked=tailored_content.get("skills_ranked"),
            ),
        }

        sections = [self._build_header()]
        seen: set = set()
        for key in order:
            if key in seen or key not in builders:
                continue
            seen.add(key)
            block = builders[key]()
            if block:
                sections.append(block)
        for key, fn in builders.items():
            if key not in seen:
                block = fn()
                if block:
                    sections.append(block)

        return sanitize_text("\n".join(sections))

    # ── Markdown section builders (deprecated path only) ──────────────────────

    def _build_header(self) -> str:
        with Session(engine) as session:
            user = session.exec(select(User).where(User.user_id == self.user_id)).first()
            if not user:
                return ""
            sh = (self._style or {}).get("header", {})
            sep = sh.get("contact_separator", " | ")
            fo  = sh.get("contact_fields", ["email", "linkedin"])
            fv  = {
                "email": user.email if user.email != "user@example.com" else None,
                "linkedin": user.linkedin_url,
                "phone": user.phone,
                "location": user.location,
                "github": f"github.com/{user.github_username}" if user.github_username else None,
            }
            contact = [fv[f] for f in fo if fv.get(f)] or [v for v in fv.values() if v]
            return "\n".join([user.name, "  ", sep.join(contact)])

    def _build_education(self) -> str:
        entries = self._get_education()
        if not entries:
            return ""
        lines = [f"{self._resolve_label('education')}  "]
        for entry in entries:
            date = self._education_date_line(entry)
            line = f"**{entry.institution}** — {self._education_degree_line(entry)}"
            if date:
                line += f", {date}"
            lines.append(line + "  ")
        return "\n".join(lines)

    def _build_projects(self, projects, label="Projects", bullet_prefix="- "):
        if not projects:
            return ""
        lines = [label]
        for proj in projects:
            lines.append(f"**{proj.get('name', 'Untitled')}**\n")
            for b in proj.get("bullets", []):
                lines.append(f"{bullet_prefix}{self._format_bullet(b)}")
            lines.append("")
        return "\n".join(lines)

    def _build_experiences(self, experiences, label="Experience", bullet_prefix="- "):
        if not experiences:
            return ""
        lines = [label]
        for exp in experiences:
            title   = exp.get("title", "")
            company = exp.get("company", "")
            start   = exp.get("start_date", "")
            end     = exp.get("end_date", "")
            lines.append(f"**{title},** {company}\t{f'{start} - {end}' if start else ''}")
            for b in exp.get("bullets", []):
                lines.append(f"{bullet_prefix}{self._format_bullet(b)}")
            lines.append("")
        return "\n".join(lines)

    def _build_skills(self, label="Skills", ranked: Optional[List[Dict]] = None):
        cats = self._get_skill_categories(ranked)
        if not cats:
            return ""
        lines = [label]
        for cat, skills in self._ordered_skill_cats(cats, ranked):
            lines.append(f"**{cat}:** {', '.join(skills)}")
        return "\n".join(lines)

    @staticmethod
    def _format_bullet(bullet: str) -> str:
        bullet = bullet.strip()
        if bullet.startswith("**"):
            return bullet
        words = bullet.split()
        if len(words) <= 3:
            return f"**{bullet}**"
        for i, word in enumerate(words):
            if i >= 4 and (word.endswith(",") or word.endswith(":")):
                lead = " ".join(words[: i + 1]).rstrip(",:")
                rest = " ".join(words[i + 1:]).lstrip(",: ")
                return f"**{lead}** {rest}" if rest else f"**{lead}**"
        lead = " ".join(words[:4])
        rest = " ".join(words[4:])
        return f"**{lead}** {rest}" if rest else f"**{lead}**"

    @staticmethod
    def _normalize_category(cat: str) -> str:
        cl = cat.lower().strip()
        if any(k in cl for k in ["language", "library", "libraries", "framework"]):
            return "Languages & Libraries"
        if any(k in cl for k in ["ml", "machine learning", "ai", "deep learning", "technique"]):
            return "AI & Machine Learning"
        if any(k in cl for k in ["data engineer", "etl", "pipeline", "database"]):
            return "Data Engineering"
        if any(k in cl for k in ["tool", "devops", "infrastructure"]):
            return "Tools"
        if any(k in cl for k in ["cloud", "aws", "gcp", "azure"]):
            return "Cloud"
        return cat
