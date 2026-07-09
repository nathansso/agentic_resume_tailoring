"""Tests for the LaTeX/docx resume formatter.

Fast tests (no pdflatex): verify .tex source structure and escaping.
Integration test (requires pdflatex): marked @pytest.mark.integration.

Test data mirrors Jake's resume sample content so the generated .tex can be
dropped directly into Overleaf against his template for visual verification.
"""
import io
import os
import shutil
import warnings

import pytest
from sqlmodel import Session

import agents.formatter as fmt_module
from agents.formatter import (
    ResumeFormatterAgent,
    _compile_tex_to_pdf,
    _convert_inline,
    _escape_tex,
    _fit_to_one_page,
    _pdf_page_count,
    _trim_one_bullet,
)
from database.models import Skill, User, UserSkill


# ── Jake-style sample data ────────────────────────────────────────────────────

_JAKE_CONTENT = {
    "experiences": [
        {
            "title": "Undergraduate Research Assistant",
            "company": "Texas A&M University",
            "start_date": "June 2020",
            "end_date": "Present",
            "location": "College Station, TX",
            "bullets": [
                "Developed a REST API using FastAPI and PostgreSQL to store data from learning management systems",
                "Developed a full-stack web application using Flask, React, PostgreSQL and Docker to analyze GitHub data",
                "Explored ways to visualize GitHub collaboration in a classroom setting",
            ],
        },
        {
            "title": "Information Technology Support Specialist",
            "company": "Southwestern University",
            "start_date": "Sep. 2018",
            "end_date": "Present",
            "location": "Georgetown, TX",
            "bullets": [
                "Communicate with managers to set up campus computers used on campus",
                "Assess and troubleshoot computer problems brought by students, faculty and staff",
                "Maintain upkeep of computers, classroom equipment, and 200 printers across campus",
            ],
        },
    ],
    "projects": [
        {
            "name": "Gitlytics",
            "tech_stack": "Python, Flask, React, PostgreSQL, Docker",
            "dates": "June 2020 -- Present",
            "bullets": [
                "Developed a full-stack web application using Flask serving a REST API with React as the frontend",
                "Implemented GitHub OAuth to get data from user's repositories",
                "Visualized GitHub data to show collaboration",
                "Used Celery and Redis for asynchronous tasks",
            ],
        },
        {
            "name": "Simple Paintball",
            "tech_stack": "Spigot API, Java, Maven, TravisCI, Git",
            "dates": "May 2018 -- May 2020",
            "bullets": [
                "Developed a Minecraft server plugin to entertain kids during free time for a previous job",
                "Published plugin to websites gaining 2K+ downloads and an average 4.5/5-star review",
                "Implemented continuous delivery using TravisCI to build the plugin upon new a release",
            ],
        },
    ],
    "skills_emphasized": ["Python", "Java", "React"],
}


def _seed_jake_user(engine) -> User:
    with Session(engine) as s:
        user = User(
            name="Jake Ryan",
            email="jake@su.edu",
            phone="123-456-7890",
            location="Georgetown, TX",
            github_username="jake",
            linkedin_url="https://linkedin.com/in/jake",
        )
        s.add(user)
        s.commit()
        s.refresh(user)

        for name, cat in [
            ("Java", "language"),
            ("Python", "language"),
            ("C/C++", "language"),
            ("SQL", "language"),
            ("JavaScript", "language"),
            ("React", "framework"),
            ("Node.js", "framework"),
            ("Flask", "framework"),
            ("FastAPI", "framework"),
            ("Git", "tool"),
            ("Docker", "tool"),
            ("TravisCI", "tool"),
        ]:
            sk = Skill(name=name, category=cat)
            s.add(sk)
            s.commit()
            s.refresh(sk)
            s.add(UserSkill(user_id=user.user_id, skill_id=sk.skill_id, confidence_score=0.9))
        s.commit()
        s.refresh(user)
        return user


# ── _escape_tex ───────────────────────────────────────────────────────────────

def test_escape_ampersand():
    assert _escape_tex("Texas A&M University") == r"Texas A\&M University"

def test_escape_percent():
    assert _escape_tex("95% coverage") == r"95\% coverage"

def test_escape_dollar():
    assert _escape_tex("$1M budget") == r"\$1M budget"

def test_escape_hash():
    assert _escape_tex("issue #42") == r"issue \#42"

def test_escape_underscore():
    assert _escape_tex("snake_case_var") == r"snake\_case\_var"

def test_escape_backslash_single_pass():
    result = _escape_tex("a\\b")
    assert result == r"a\textbackslash{}b"
    assert result.count(r"\textbackslash{}") == 1

def test_escape_braces():
    assert _escape_tex("{hello}") == r"\{hello\}"

def test_escape_tilde():
    assert _escape_tex("~") == r"\textasciitilde{}"


# ── _convert_inline ───────────────────────────────────────────────────────────

def test_convert_bold_to_textbf():
    result = _convert_inline("**Developed** a REST API")
    assert r"\textbf{Developed}" in result
    assert "**" not in result

def test_convert_italic_to_textit():
    result = _convert_inline("based on *The Legend of Zelda*")
    assert r"\textit{The Legend of Zelda}" in result

def test_convert_escapes_special_chars_outside_markers():
    result = _convert_inline("Contributed 50% via Git & GitHub")
    assert r"50\%" in result
    assert r"Git \& GitHub" in result

def test_convert_bold_content_with_special_chars():
    result = _convert_inline("**Python & Go**")
    assert r"\textbf{Python \& Go}" in result

def test_convert_no_markers_still_escapes():
    result = _convert_inline("2K+ downloads at 4.5/5 stars")
    assert "**" not in result
    assert result == r"2K+ downloads at 4.5/5 stars"

def test_convert_markdown_link_to_href():
    result = _convert_inline("Shipped a [live demo](https://example.com/demo) for the app")
    assert r"\href{https://example.com/demo}{\underline{live demo}}" in result
    assert "[" not in result and "]" not in result

def test_convert_markdown_link_escapes_special_chars():
    result = _convert_inline("See [repo](https://example.com/a_b?x=1&y=2)")
    assert r"\underline{repo}" in result
    assert r"a\_b" in result


# ── .tex document structure ───────────────────────────────────────────────────

def test_tex_documentclass(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(_JAKE_CONTENT)
    assert r"\documentclass[letterpaper,11pt]{article}" in tex

def test_tex_jake_preamble_packages(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(_JAKE_CONTENT)
    for pkg in [r"\usepackage{marvosym}", r"\usepackage{latexsym}",
                r"\usepackage{enumitem}", r"\usepackage[hidelinks]{hyperref}",
                r"\usepackage{fancyhdr}", r"\usepackage{tabularx}",
                r"\usepackage{titlesec}", r"\pdfgentounicode=1"]:
        assert pkg in tex, f"Missing package: {pkg}"

def test_tex_jake_custom_commands(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(_JAKE_CONTENT)
    for cmd in [r"\resumeItem", r"\resumeSubheading", r"\resumeSubSubheading",
                r"\resumeProjectHeading", r"\resumeSubItem",
                r"\resumeSubHeadingListStart", r"\resumeSubHeadingListEnd",
                r"\resumeItemListStart", r"\resumeItemListEnd"]:
        assert cmd in tex, f"Missing command definition: {cmd}"

def test_tex_document_begin_end(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(_JAKE_CONTENT)
    assert r"\begin{document}" in tex
    assert r"\end{document}" in tex

def test_tex_header_name_and_contact(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(_JAKE_CONTENT)
    assert r"\textbf{\Huge \scshape Jake Ryan}" in tex
    assert "123-456-7890" in tex
    assert r"mailto:jake@su.edu" in tex
    assert "linkedin.com/in/jake" in tex
    assert "github.com/jake" in tex

def test_tex_header_includes_portfolio_url(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    with Session(isolated_engine) as s:
        db_user = s.get(User, user.user_id)
        db_user.portfolio_url = "https://jakeryan.dev"
        s.add(db_user)
        s.commit()
    tex = ResumeFormatterAgent(user.user_id)._build_tex(_JAKE_CONTENT)
    assert r"\href{https://jakeryan.dev}" in tex
    assert "jakeryan.dev" in tex

def test_tex_experience_uses_resumesubheading(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(_JAKE_CONTENT)
    # arg1=title, arg2=dates, arg3=company, arg4=location (Jake's order)
    assert "{Undergraduate Research Assistant}{June 2020 -- Present}" in tex
    assert "{Texas A\\&M University}{College Station, TX}" in tex

def test_tex_experience_bullets_use_resumeitem(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(_JAKE_CONTENT)
    assert r"\resumeItemListStart" in tex
    assert r"\resumeItemListEnd" in tex
    assert r"\resumeItem{Developed a REST API" in tex

def test_tex_projects_use_resumeprojectheading(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(_JAKE_CONTENT)
    assert r"\resumeProjectHeading" in tex
    assert r"\textbf{Gitlytics} $|$ \emph{Python, Flask, React, PostgreSQL, Docker}" in tex
    assert "{June 2020 -- Present}" in tex

def test_tex_projects_render_repo_and_demo_links(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    content = {
        **_JAKE_CONTENT,
        "projects": [
            {
                **_JAKE_CONTENT["projects"][0],
                "repo_url": "https://github.com/jake/gitlytics",
                "demo_url": "https://gitlytics.example.com",
            },
        ],
    }
    tex = ResumeFormatterAgent(user.user_id)._build_tex(content)
    assert r"\href{https://github.com/jake/gitlytics}{\underline{Repo}}" in tex
    assert r"\href{https://gitlytics.example.com}{\underline{Demo}}" in tex

def test_tex_projects_omit_links_when_absent(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(_JAKE_CONTENT)
    assert r"\underline{Repo}" not in tex
    assert r"\underline{Demo}" not in tex


def test_tex_skills_section_is_technical_skills(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(_JAKE_CONTENT)
    assert r"\section{Technical Skills}" in tex
    assert r"\textbf{Languages \& Libraries}" in tex

def test_tex_section_order_respected(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(
        _JAKE_CONTENT, section_order=["projects", "experience", "skills", "education"]
    )
    assert tex.index(r"\section{Projects}") < tex.index(r"\section{Experience}")

def test_tex_header_stays_above_sections_with_custom_order(isolated_engine, monkeypatch):
    """Name/contact header is pinned at the top regardless of section_order (issue 22)."""
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(
        _JAKE_CONTENT, section_order=["projects", "skills", "experience", "education"]
    )
    first_section = tex.index(r"\section{")
    assert tex.index("Jake Ryan") < first_section
    assert tex.index("123-456-7890") < first_section

def test_tex_section_markers_present(isolated_engine, monkeypatch):
    """Every emitted block carries an ART-SECTION marker (issue #71) so the web
    editor can reorder sections/bullets as text-block moves."""
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(_JAKE_CONTENT)
    assert "%% ART-SECTION: header" in tex
    for key in ("experience", "projects", "skills"):
        assert f"%% ART-SECTION: {key}" in tex, f"Missing marker for {key}"
    # Marker sits directly above the block it labels
    assert tex.index("%% ART-SECTION: experience") < tex.index(r"\section{Experience}")
    assert tex.index(r"\section{Experience}") < tex.index("%% ART-SECTION: projects")


def test_tex_section_markers_follow_custom_order(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(
        _JAKE_CONTENT, section_order=["projects", "experience", "skills", "education"]
    )
    assert tex.index("%% ART-SECTION: header") < tex.index("%% ART-SECTION: projects")
    assert tex.index("%% ART-SECTION: projects") < tex.index("%% ART-SECTION: experience")
    assert tex.index("%% ART-SECTION: experience") < tex.index("%% ART-SECTION: skills")


def test_tex_special_chars_escaped(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    content = {
        "experiences": [{
            "title": "Engineer & Lead",
            "company": "50% Corp",
            "start_date": "Jan 2023",
            "end_date": "Present",
            "location": "",
            "bullets": ["Managed $1M budget across 3 accounts"],
        }],
        "projects": [],
        "skills_emphasized": [],
    }
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id)._build_tex(content)
    assert r"Engineer \& Lead" in tex
    assert r"50\% Corp" in tex
    assert r"\$1M budget" in tex


# ── zero-bullet entries and placeholder dates (issue #72) ─────────────────────

def _one_exp_content(**overrides):
    exp = {
        "title": "Analyst", "company": "Acme", "start_date": "2023-01",
        "end_date": "Present", "location": "", "bullets": ["Did a thing"],
    }
    exp.update(overrides)
    return {"experiences": [exp], "projects": [], "skills_emphasized": []}


def _document_body(tex: str) -> str:
    """The part after \\begin{document} — excludes the preamble's \\newcommand defs."""
    return tex.split(r"\begin{document}", 1)[1]


def test_experience_with_no_bullets_skips_itemize(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    # A description-only experience still renders its heading but no empty itemize.
    content = _one_exp_content(bullets=[])
    body = _document_body(ResumeFormatterAgent(user.user_id)._build_tex(content))
    assert "{Analyst}" in body
    assert r"\resumeItemListStart" not in body
    assert r"\resumeItemListEnd" not in body


def test_project_with_no_bullets_skips_itemize(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    content = {
        "experiences": [],
        "projects": [{"name": "Solo", "bullets": []}],
        "skills_emphasized": [],
    }
    body = _document_body(ResumeFormatterAgent(user.user_id)._build_tex(content))
    assert r"\textbf{Solo}" in body
    assert r"\resumeItemListStart" not in body


def test_placeholder_start_date_is_omitted(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    content = _one_exp_content(start_date="Not specified", end_date="Present")
    tex = ResumeFormatterAgent(user.user_id)._build_tex(content)
    assert "Not specified" not in tex
    # Only the real end date renders, with no dangling "--" separator.
    assert "{Analyst}{Present}" in tex


def test_placeholder_both_dates_render_blank(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    content = _one_exp_content(start_date="unknown", end_date="N/A")
    tex = ResumeFormatterAgent(user.user_id)._build_tex(content)
    assert "unknown" not in tex and "N/A" not in tex
    assert "{Analyst}{}" in tex


# ── format_tex() ──────────────────────────────────────────────────────────────

def test_format_tex_returns_string(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id).format_tex(_JAKE_CONTENT)
    assert isinstance(tex, str)
    assert r"\documentclass" in tex


# ── format_docx() ─────────────────────────────────────────────────────────────

def test_format_docx_returns_bytes(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    result = ResumeFormatterAgent(user.user_id).format_docx(_JAKE_CONTENT)
    assert isinstance(result, bytes)
    assert len(result) > 500

def test_format_docx_includes_project_links(isolated_engine, monkeypatch):
    from docx import Document
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    content = {
        **_JAKE_CONTENT,
        "projects": [
            {**_JAKE_CONTENT["projects"][0], "repo_url": "https://github.com/jake/gitlytics"},
        ],
    }
    docx_bytes = ResumeFormatterAgent(user.user_id).format_docx(content)
    full_text = "\n".join(p.text for p in Document(io.BytesIO(docx_bytes)).paragraphs)
    assert "https://github.com/jake/gitlytics" in full_text

def test_format_docx_is_valid_docx(isolated_engine, monkeypatch):
    """Verify the bytes can be opened as a Word document."""
    from docx import Document
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    docx_bytes = ResumeFormatterAgent(user.user_id).format_docx(_JAKE_CONTENT)
    doc = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Jake Ryan" in full_text
    assert "Gitlytics" in full_text


# ── format_markdown() deprecation ────────────────────────────────────────────

def test_format_markdown_emits_deprecation_warning(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        ResumeFormatterAgent(user.user_id).format_markdown(_JAKE_CONTENT)
    assert any(issubclass(w.category, DeprecationWarning) for w in caught)


# ── save .tex to Downloads for manual Overleaf verification ──────────────────

def test_save_tex_to_downloads(isolated_engine, monkeypatch):
    """Generate a Jake-layout .tex and save to ~/Downloads for Overleaf inspection."""
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    tex = ResumeFormatterAgent(user.user_id).format_tex(_JAKE_CONTENT)

    downloads = os.path.join(os.path.expanduser("~"), "Downloads", "jake_resume_art.tex")
    with open(downloads, "w", encoding="utf-8") as f:
        f.write(tex)

    assert os.path.exists(downloads)
    assert os.path.getsize(downloads) > 1000


# ── integration: full pdflatex compile ───────────────────────────────────────

@pytest.mark.integration
def test_pdflatex_produces_valid_pdf(isolated_engine, monkeypatch):
    """Requires pdflatex. Verifies the .tex compiles and output is a valid PDF."""
    if not shutil.which("pdflatex"):
        pytest.skip("pdflatex not installed")

    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    pdf_bytes = ResumeFormatterAgent(user.user_id).format_pdf(_JAKE_CONTENT)

    assert isinstance(pdf_bytes, bytes)
    assert len(pdf_bytes) > 1000
    assert pdf_bytes[:4] == b"%PDF"

    # Also save PDF to Downloads for visual inspection
    downloads = os.path.join(os.path.expanduser("~"), "Downloads", "jake_resume_art.pdf")
    with open(downloads, "wb") as f:
        f.write(pdf_bytes)
    assert os.path.exists(downloads)


# ── one-page fit enforcement (issue #34), pure logic — no LaTeX engine ─────────

def test_trim_one_bullet_trims_project_fat_to_floor_first():
    # Above the floor: the lowest-ranked project's last bullet goes first.
    content = {
        "projects": [
            {"name": "A", "bullets": ["a1", "a2", "a3"]},
            {"name": "B", "bullets": ["b1", "b2", "b3"]},
        ],
    }
    trimmed = _trim_one_bullet(content)
    assert trimmed["projects"][1]["bullets"] == ["b1", "b2"]   # B trimmed toward floor
    assert trimmed["projects"][0]["bullets"] == ["a1", "a2", "a3"]
    # Original is untouched (operates on a copy).
    assert content["projects"][1]["bullets"] == ["b1", "b2", "b3"]


def test_trim_one_bullet_drops_whole_project_at_floor():
    # Both projects already at the floor (2): drop the weakest entire project
    # rather than starve the survivors (issue #72).
    content = {
        "projects": [
            {"name": "A", "bullets": ["a1", "a2"]},
            {"name": "B", "bullets": ["b1", "b2"]},
        ],
    }
    trimmed = _trim_one_bullet(content)
    assert [p["name"] for p in trimmed["projects"]] == ["A"]
    assert trimmed["projects"][0]["bullets"] == ["a1", "a2"]  # survivor keeps its bullets


def test_trim_one_bullet_never_empties_a_project():
    content = {"projects": [{"name": "A", "bullets": ["only"]}]}
    # A single-bullet lone project can't be trimmed; with no other text, None.
    assert _trim_one_bullet(content) is None


def test_trim_one_bullet_falls_back_to_experiences():
    # One lone project already at one bullet, experiences carry the fat.
    content = {
        "projects": [{"name": "A", "bullets": ["a1"]}],
        "experiences": [{"title": "T", "bullets": ["e1", "e2", "e3"]}],
    }
    trimmed = _trim_one_bullet(content)
    assert trimmed["experiences"][0]["bullets"] == ["e1", "e2"]  # exp fat to floor
    assert trimmed["projects"][0]["bullets"] == ["a1"]


def test_trim_one_bullet_protects_experiences_below_floor_last():
    # Single project pinned at one bullet, experience at the floor: the last
    # resort shaves the experience below the floor rather than failing.
    content = {
        "projects": [{"name": "A", "bullets": ["a1"]}],
        "experiences": [{"title": "T", "bullets": ["e1", "e2"]}],
    }
    trimmed = _trim_one_bullet(content)
    assert trimmed["experiences"][0]["bullets"] == ["e1"]


def test_trim_one_bullet_exhausted_returns_none():
    content = {
        "projects": [{"name": "A", "bullets": ["a"]}],
        "experiences": [{"title": "T", "bullets": ["e"]}],
    }
    assert _trim_one_bullet(content) is None


def test_fit_to_one_page_trims_until_it_fits():
    # Fake render/page-count: "fits" once total bullets <= 4.
    def total_bullets(c):
        return sum(len(i["bullets"]) for i in c.get("projects", []))

    def page_count(rendered):
        return 1 if rendered <= 4 else 2

    content = {
        "projects": [
            {"name": "A", "bullets": ["a1", "a2", "a3"]},
            {"name": "B", "bullets": ["b1", "b2", "b3"]},
        ],
    }
    fitted = _fit_to_one_page(content, total_bullets, page_count)
    assert total_bullets(fitted) <= 4
    # Fat was trimmed from the lowest project (B) first, keeping both at/above floor.
    assert all(len(p["bullets"]) >= 2 for p in fitted["projects"])
    assert len(fitted["projects"]) == 2


def test_fit_to_one_page_stops_when_exhausted():
    # Never fits, but the reducer bottoms out — loop must terminate, not hang.
    content = {"projects": [{"name": "A", "bullets": ["a1", "a2"]}]}
    fitted = _fit_to_one_page(content, lambda c: 0, lambda r: 2)
    assert fitted["projects"][0]["bullets"] == ["a1"]  # single project shaved to one, then stopped


def test_fit_content_to_one_page_trims_overflow(isolated_engine, monkeypatch):
    """fit_content_to_one_page trims via the shared reducer using a real tex
    build but a stubbed compile/page-count (no LaTeX engine needed)."""
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    agent = ResumeFormatterAgent(user.user_id)

    def bullets_total(content):
        return sum(len(i.get("bullets", [])) for i in content.get("projects", []))

    original_total = bullets_total(_JAKE_CONTENT)
    compiled = {}

    def fake_compile(tex_str):
        return tex_str.encode()

    def fake_page_count(pdf_bytes):
        # "Fits" once at least two bullets have been trimmed.
        n = pdf_bytes.decode().count(r"\resumeItem{")
        compiled["last_items"] = n
        return 1 if n <= compiled["first_items"] - 2 else 2

    monkeypatch.setattr(fmt_module, "_compile_tex_to_pdf", fake_compile)
    first_tex = agent.format_tex(_JAKE_CONTENT)
    compiled["first_items"] = first_tex.count(r"\resumeItem{")
    monkeypatch.setattr(fmt_module, "_pdf_page_count", fake_page_count)

    fitted = agent.fit_content_to_one_page(_JAKE_CONTENT)
    assert bullets_total(fitted) < original_total
    # Original content object untouched (trim works on copies).
    assert bullets_total(_JAKE_CONTENT) == original_total


def test_fit_content_to_one_page_no_engine_returns_unchanged(isolated_engine, monkeypatch):
    """Offline/test environments without a LaTeX engine get the content back
    unchanged instead of an exception."""
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    agent = ResumeFormatterAgent(user.user_id)

    def raising_compile(tex_str):
        raise RuntimeError("No LaTeX engine found.")

    monkeypatch.setattr(fmt_module, "_compile_tex_to_pdf", raising_compile)
    assert agent.fit_content_to_one_page(_JAKE_CONTENT) is _JAKE_CONTENT


def test_fit_enforcement_keeps_preamble_byte_identical(isolated_engine, monkeypatch):
    """Trimming text must not alter the preamble/geometry/font (issue #34)."""
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    agent = ResumeFormatterAgent(user.user_id)

    trimmed_content = _trim_one_bullet(_JAKE_CONTENT)
    assert trimmed_content is not None

    baseline_tex = agent.format_tex(_JAKE_CONTENT)
    trimmed_tex = agent.format_tex(trimmed_content)

    def preamble(tex):
        return tex.split(r"\begin{document}")[0]

    assert preamble(baseline_tex) == preamble(trimmed_tex)


# ── one-page fit enforcement (issue #34), integration — requires a LaTeX engine ─

# A deliberately long resume: 4 dense projects + 3 experiences overflow one page.
_OVERFLOW_CONTENT = {
    "experiences": [
        {
            "title": f"Senior Engineer {n}",
            "company": f"Company {n}",
            "start_date": "Jan 2020",
            "end_date": "Present",
            "location": "Remote",
            "bullets": [
                "Built scalable distributed services handling millions of requests per day with python and go",
                "Led a team designing fault-tolerant data pipelines on kubernetes with terraform and airflow",
                "Reduced infrastructure cost 40% by re-architecting the streaming layer around kafka and redis",
                "Mentored five engineers and drove adoption of rigorous integration testing and observability",
            ],
        }
        for n in range(3)
    ],
    "projects": [
        {
            "name": f"Project {n}",
            "tech_stack": "Python, FastAPI, React, PostgreSQL, Docker, Kubernetes",
            "dates": "2021 -- Present",
            "bullets": [
                "Developed a full-stack platform with fastapi and react serving thousands of users",
                "Implemented oauth, role-based access control, and audit logging across the system",
                "Containerized the deployment with docker and kubernetes and automated ci/cd pipelines",
                "Instrumented grafana dashboards and structured logging for production observability",
            ],
        }
        for n in range(4)
    ],
    "skills_emphasized": ["Python", "React", "Kubernetes"],
}


def _no_latex_engine() -> bool:
    return not shutil.which("tectonic") and not shutil.which("pdflatex") and not any(
        os.path.exists(os.path.join(os.path.dirname(__import__("sys").executable), exe))
        for exe in ("tectonic.exe", "tectonic")
    )


@pytest.mark.integration
def test_format_pdf_overflow_fits_one_page(isolated_engine, monkeypatch):
    """A long 4-project resume is trimmed down to exactly one page (k=4 case)."""
    if _no_latex_engine():
        pytest.skip("no LaTeX engine (tectonic/pdflatex) installed")

    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    pdf_bytes = ResumeFormatterAgent(user.user_id).format_pdf(_OVERFLOW_CONTENT)

    assert pdf_bytes[:4] == b"%PDF"
    assert _pdf_page_count(pdf_bytes) == 1


@pytest.mark.integration
def test_format_pdf_short_resume_stays_one_page(isolated_engine, monkeypatch):
    """A 2-project resume already fits and is returned as one page (k=2 case)."""
    if _no_latex_engine():
        pytest.skip("no LaTeX engine (tectonic/pdflatex) installed")

    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _seed_jake_user(isolated_engine)
    pdf_bytes = ResumeFormatterAgent(user.user_id).format_pdf(_JAKE_CONTENT)

    assert _pdf_page_count(pdf_bytes) == 1
