"""Web ingestion regression tests.

1. Resume ingestion must work without docling installed — the production
   Docker image excludes it because docling pulls in PyTorch and loading its
   layout models OOM-kills the 512 MB Fly.io VM.
2. /api/jobs must resolve with and without a trailing slash — the SPA
   catch-all route swallows FastAPI's automatic slash-redirect, which turned
   the slash-less form into a hard 404 in production.
"""
import sys

import pytest


@pytest.fixture()
def no_docling(monkeypatch):
    """Simulate docling being absent even when it is installed locally."""
    monkeypatch.setitem(sys.modules, "docling", None)
    monkeypatch.setitem(sys.modules, "docling.document_converter", None)


def _make_docx(path):
    import docx
    document = docx.Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("jane@example.com | linkedin.com/in/janedoe")
    document.add_paragraph("Skills")
    document.add_paragraph("Python, SQL")
    document.add_paragraph("Experience")
    document.add_paragraph("Software Engineer at Acme")
    document.save(str(path))


def test_resume_ingest_falls_back_without_docling(no_docling, tmp_path):
    from ingestion.resume import ResumeIngestor

    docx_path = tmp_path / "resume.docx"
    _make_docx(docx_path)

    data = ResumeIngestor().ingest(str(docx_path))

    assert "Jane Doe" in data["full_text"]
    assert "skills" in data["parsed_sections"]
    assert any("Python" in item["text"] for item in data["parsed_sections"]["skills"])
    assert "experience" in data["parsed_sections"]
    assert data["resume_style"]["section_order"] == ["skills", "experience"]


def test_plain_text_extraction_segments_sections(no_docling, tmp_path):
    from ingestion.resume import ResumeIngestor

    txt_path = tmp_path / "resume.txt"
    txt_path.write_text(
        "John Smith\n\nEducation\nBS Computer Science\n\nProjects\nBuilt a web app\n",
        encoding="utf-8",
    )

    data = ResumeIngestor().ingest(str(txt_path))

    assert "education" in data["parsed_sections"]
    assert "projects" in data["parsed_sections"]
    assert any("BS Computer Science" in i["text"] for i in data["parsed_sections"]["education"])


def test_pypdf_extraction_handles_pdf_without_docling(no_docling, tmp_path):
    from pypdf import PdfWriter

    from ingestion.document_text import extract_text_lightweight

    pdf_path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with open(pdf_path, "wb") as fh:
        writer.write(fh)

    # A blank page has no text — the point is that extraction runs the pypdf
    # path without docling and without raising.
    assert extract_text_lightweight(str(pdf_path)) == ""


def test_linkedin_pdf_ingest_uses_fallback_without_docling(no_docling, tmp_path):
    from ingestion.linkedin import LinkedInIngestor

    txt_path = tmp_path / "profile.txt"
    txt_path.write_text("Jane Doe\nSoftware Engineer", encoding="utf-8")

    data = LinkedInIngestor().ingest_pdf(str(txt_path))

    assert data["source_type"] == "linkedin"
    assert "Jane Doe" in data["full_text"]


def test_jobs_collection_route_resolves_without_trailing_slash():
    from fastapi.testclient import TestClient

    from web.app import create_app

    client = TestClient(create_app())
    resp = client.get("/api/jobs")
    # 401 (auth required) proves the route resolved; the regression was a 404
    # from the SPA catch-all.
    assert resp.status_code == 401
    assert client.get("/api/jobs/").status_code == 401
