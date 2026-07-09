"""Per-user education storage and rendering (issue #73).

The formatter previously hardcoded one user's education into every export,
leaking it to all users. These tests pin the fix: education comes from the
per-user Education table, and a user with no rows gets no education section.
"""
import io

import pytest
from sqlmodel import Session, select

import agents.formatter as fmt_module
import agents.parser as parser_module
from agents.formatter import ResumeFormatterAgent
from database.models import Education, User

# The strings that used to be hardcoded — they must never appear for a user
# who does not have them in the DB.
_LEAKED_MARKERS = [
    "University of California, San Diego",
    "M.S. Data Science",
    "Mathematics",
]


def _make_user(engine, email="edu@example.com", name="Edu User") -> User:
    with Session(engine) as s:
        user = User(name=name, email=email)
        s.add(user)
        s.commit()
        s.refresh(user)
        return user


def _seed_education(engine, user_id, entries):
    with Session(engine) as s:
        for e in entries:
            s.add(Education(user_id=user_id, **e))
        s.commit()


def _make_parser(isolated_engine, monkeypatch, user):
    """ResumeParserAgent wired to the test DB, bypassing __init__ (no LLM)."""
    monkeypatch.setattr(parser_module, "engine", isolated_engine)
    agent = parser_module.ResumeParserAgent.__new__(parser_module.ResumeParserAgent)
    agent.user = user
    agent.llm = None
    return agent


_ENTRIES = [
    {
        "institution": "Texas A&M University",
        "degree": "B.S. Computer Science",
        "location": "College Station, TX",
        "end_date": "May 2021",
        "gpa": "3.8/4.0",
    },
    {
        "institution": "Southwestern University",
        "degree": "A.A. Liberal Arts",
        "location": "Georgetown, TX",
        "start_date": "Aug 2015",
        "end_date": "May 2017",
    },
]


# ── LaTeX path ────────────────────────────────────────────────────────────────

def test_tex_education_renders_user_rows(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _make_user(isolated_engine)
    _seed_education(isolated_engine, user.user_id, _ENTRIES)

    tex = ResumeFormatterAgent(user.user_id)._build_tex({})

    assert r"Texas A\&M University" in tex
    assert "B.S. Computer Science, GPA: 3.8/4.0" in tex
    assert "Southwestern University" in tex
    assert "Aug 2015 -- May 2017" in tex
    for marker in _LEAKED_MARKERS:
        assert marker not in tex


def test_tex_education_omitted_when_no_rows(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _make_user(isolated_engine)

    tex = ResumeFormatterAgent(user.user_id)._build_tex({})

    assert r"\section{Education}" not in tex
    for marker in _LEAKED_MARKERS:
        assert marker not in tex


def test_tex_education_isolated_between_users(isolated_engine, monkeypatch):
    """User B's export must not contain user A's education."""
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user_a = _make_user(isolated_engine, email="a@example.com", name="User A")
    user_b = _make_user(isolated_engine, email="b@example.com", name="User B")
    _seed_education(isolated_engine, user_a.user_id, [_ENTRIES[0]])
    _seed_education(
        isolated_engine, user_b.user_id,
        [{"institution": "MIT", "degree": "B.S. Physics", "end_date": "June 2020"}],
    )

    tex_b = ResumeFormatterAgent(user_b.user_id)._build_tex({})

    assert "MIT" in tex_b
    assert "Texas" not in tex_b


# ── DOCX path ─────────────────────────────────────────────────────────────────

def test_docx_education_renders_user_rows(isolated_engine, monkeypatch):
    from docx import Document

    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _make_user(isolated_engine)
    _seed_education(isolated_engine, user.user_id, [_ENTRIES[0]])

    docx_bytes = ResumeFormatterAgent(user.user_id).format_docx({})
    text = "\n".join(p.text for p in Document(io.BytesIO(docx_bytes)).paragraphs)

    assert "Texas A&M University" in text
    assert "B.S. Computer Science, GPA: 3.8/4.0" in text
    for marker in _LEAKED_MARKERS:
        assert marker not in text


def test_docx_education_omitted_when_no_rows(isolated_engine, monkeypatch):
    from docx import Document

    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _make_user(isolated_engine)

    docx_bytes = ResumeFormatterAgent(user.user_id).format_docx({})
    text = "\n".join(p.text for p in Document(io.BytesIO(docx_bytes)).paragraphs)

    assert "EDUCATION" not in text.upper()
    for marker in _LEAKED_MARKERS:
        assert marker not in text


# ── Markdown (deprecated) path ────────────────────────────────────────────────

def test_markdown_education_from_db(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _make_user(isolated_engine)
    _seed_education(isolated_engine, user.user_id, [_ENTRIES[0]])

    block = ResumeFormatterAgent(user.user_id)._build_education()

    assert "**Texas A&M University**" in block
    assert "B.S. Computer Science" in block
    for marker in _LEAKED_MARKERS:
        assert marker not in block


def test_markdown_education_empty_when_no_rows(isolated_engine, monkeypatch):
    monkeypatch.setattr(fmt_module, "engine", isolated_engine)
    user = _make_user(isolated_engine)
    assert ResumeFormatterAgent(user.user_id)._build_education() == ""


# ── Parser: resume extraction save + dedup ────────────────────────────────────

def test_save_education_persists_rows(isolated_engine, monkeypatch):
    user = _make_user(isolated_engine)
    agent = _make_parser(isolated_engine, monkeypatch, user)

    agent._save_education(
        [{"institution": "Texas A&M University", "degree": "B.S. Computer Science",
          "location": "College Station, TX", "end_date": "May 2021", "gpa": 3.8}],
        "resume",
    )

    with Session(isolated_engine) as s:
        rows = s.exec(select(Education).where(Education.user_id == user.user_id)).all()
    assert len(rows) == 1
    assert rows[0].institution == "Texas A&M University"
    assert rows[0].gpa == "3.8"  # numeric GPA coerced to string


def test_save_education_dedups_on_reingest(isolated_engine, monkeypatch):
    user = _make_user(isolated_engine)
    agent = _make_parser(isolated_engine, monkeypatch, user)
    data = [{"institution": "Texas A&M University", "degree": "B.S. Computer Science"}]

    agent._save_education(data, "resume")
    agent._save_education(data, "resume")

    with Session(isolated_engine) as s:
        rows = s.exec(select(Education).where(Education.user_id == user.user_id)).all()
    assert len(rows) == 1


def test_save_education_skips_blank_institution(isolated_engine, monkeypatch):
    user = _make_user(isolated_engine)
    agent = _make_parser(isolated_engine, monkeypatch, user)

    agent._save_education([{"institution": "  ", "degree": "B.S."}], "resume")

    with Session(isolated_engine) as s:
        rows = s.exec(select(Education).where(Education.user_id == user.user_id)).all()
    assert rows == []


def test_save_education_keeps_distinct_degree_levels(isolated_engine, monkeypatch):
    """The M.S. and B.S. at one school are distinct rows, never collapsed."""
    user = _make_user(isolated_engine)
    agent = _make_parser(isolated_engine, monkeypatch, user)
    agent._save_education(
        [{"institution": "University of California, San Diego", "degree": "M.S. Data Science"},
         {"institution": "University of California, San Diego",
          "degree": "B.S. Mathematics & Economics, Minor in Data Science"}],
        "resume",
    )
    with Session(isolated_engine) as s:
        rows = s.exec(select(Education).where(Education.user_id == user.user_id)).all()
    assert len(rows) == 2


def test_save_education_merges_same_level_duplicate(isolated_engine, monkeypatch):
    """A re-emitted undergrad row with a trimmed degree string (same level) is
    merged, not duplicated — the double-undergrad bug."""
    user = _make_user(isolated_engine)
    agent = _make_parser(isolated_engine, monkeypatch, user)
    agent._save_education(
        [{"institution": "University of California, San Diego",
          "degree": "B.S. Mathematics & Economics, Minor in Data Science",
          "end_date": "June 2025"}],
        "resume",
    )
    # Second ingest: same undergrad, degree captured more tersely, missing date.
    agent._save_education(
        [{"institution": "University of California, San Diego",
          "degree": "B.S. Mathematics & Economics"}],
        "resume",
    )
    with Session(isolated_engine) as s:
        rows = s.exec(select(Education).where(Education.user_id == user.user_id)).all()
    assert len(rows) == 1
    assert rows[0].end_date == "June 2025"  # richer row's data preserved


def test_heal_education_merges_existing_double_undergrad(isolated_engine, monkeypatch):
    """Self-heal collapses two persisted same-level rows while keeping the M.S.
    distinct — cleans up dirty DBs on next ingest without a hard reset."""
    user = _make_user(isolated_engine)
    _seed_education(
        isolated_engine, user.user_id,
        [{"institution": "University of California, San Diego", "degree": "M.S. Data Science",
          "gpa": "4.0/4.0"},
         {"institution": "University of California, San Diego",
          "degree": "B.S. Mathematics & Economics, Minor in Data Science",
          "end_date": "June 2025"},
         {"institution": "University of California, San Diego",
          "degree": "B.S. Mathematics & Economics"}],
    )
    with Session(isolated_engine) as s:
        removed = parser_module.ResumeParserAgent._heal_education(s, user.user_id)
        s.commit()
    assert removed == 1
    with Session(isolated_engine) as s:
        rows = s.exec(select(Education).where(Education.user_id == user.user_id)).all()
    degrees = sorted(r.degree for r in rows)
    assert len(rows) == 2  # M.S. + one merged B.S.
    assert any(d.startswith("M.S.") for d in degrees)
    assert any(d.startswith("B.S.") for d in degrees)


def test_heal_education_idempotent_on_clean_data(isolated_engine, monkeypatch):
    user = _make_user(isolated_engine)
    _seed_education(
        isolated_engine, user.user_id,
        [{"institution": "UCSD", "degree": "M.S. Data Science"},
         {"institution": "UCSD", "degree": "B.S. Mathematics"}],
    )
    with Session(isolated_engine) as s:
        assert parser_module.ResumeParserAgent._heal_education(s, user.user_id) == 0
        s.commit()
    with Session(isolated_engine) as s:
        rows = s.exec(select(Education).where(Education.user_id == user.user_id)).all()
    assert len(rows) == 2


# ── Service + API: education visible in the Data Explorer ────────────────────

def test_get_education_service_shape(isolated_engine):
    import services as services_module
    user = _make_user(isolated_engine)
    _seed_education(isolated_engine, user.user_id, [_ENTRIES[0]])

    rows = services_module.get_education(user.user_id)

    assert rows == [{
        "institution": "Texas A&M University",
        "degree": "B.S. Computer Science",
        "location": "College Station, TX",
        "start": "",
        "end": "May 2021",
        "gpa": "3.8/4.0",
    }]
    assert services_module.get_education(None) == []


def test_education_endpoint_scoped_to_caller(isolated_engine, monkeypatch):
    import database.db as db_module
    monkeypatch.setattr(db_module, "engine", isolated_engine)
    import services as services_module
    monkeypatch.setattr(services_module, "engine", isolated_engine)

    from fastapi.testclient import TestClient
    from web.app import create_app
    import web.auth as web_auth_module

    alice = _make_user(isolated_engine, email="edu-a@example.com", name="A")
    bob = _make_user(isolated_engine, email="edu-b@example.com", name="B")
    _seed_education(isolated_engine, alice.user_id, [_ENTRIES[0]])

    def client_for(user):
        app = create_app()
        app.dependency_overrides[web_auth_module.get_current_user] = lambda: user
        return TestClient(app, raise_server_exceptions=True)

    rows_alice = client_for(alice).get("/api/profile/education").json()
    rows_bob = client_for(bob).get("/api/profile/education").json()

    assert [r["institution"] for r in rows_alice] == ["Texas A&M University"]
    assert rows_bob == []


# ── Parser: LinkedIn structured mapping ───────────────────────────────────────

_LINKEDIN_RECORD = {
    "education": [
        {"title": "UCSD", "degree": "BS", "field": "CS",
         "start_year": "2019", "end_year": "2023"},
    ],
}


def test_linkedin_structured_education_saved(isolated_engine, monkeypatch):
    user = _make_user(isolated_engine)
    agent = _make_parser(isolated_engine, monkeypatch, user)

    agent._save_linkedin_structured(_LINKEDIN_RECORD)

    with Session(isolated_engine) as s:
        rows = s.exec(select(Education).where(Education.user_id == user.user_id)).all()
    assert len(rows) == 1
    assert rows[0].institution == "UCSD"
    assert rows[0].degree == "BS, CS"
    assert rows[0].start_date == "2019"
    assert rows[0].end_date == "2023"


def test_linkedin_education_merges_with_existing(isolated_engine, monkeypatch):
    """A resume-ingested institution is not duplicated by the LinkedIn record."""
    user = _make_user(isolated_engine)
    _seed_education(
        isolated_engine, user.user_id,
        [{"institution": "UCSD", "degree": "B.S. Computer Science"}],
    )
    agent = _make_parser(isolated_engine, monkeypatch, user)

    agent._save_linkedin_structured(_LINKEDIN_RECORD)
    agent._save_linkedin_structured(_LINKEDIN_RECORD)  # re-ingest is idempotent too

    with Session(isolated_engine) as s:
        rows = s.exec(select(Education).where(Education.user_id == user.user_id)).all()
    assert len(rows) == 1
    assert rows[0].degree == "B.S. Computer Science"
